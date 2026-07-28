"""
VERSE Parsers - Document Text & Table Extractor Engine (PDF, DOCX, TXT).
"""

import os
from typing import List
import fitz  # PyMuPDF
from docx import Document

from app.core.errors import ParsingError
from app.core.logging import get_logger

logger = get_logger(__name__)


def extract_text(file_path: str) -> str:
    """
    Extract plain text from a screenplay or production document on disk.
    
    Supports: .pdf, .docx, .txt
    """
    if not os.path.exists(file_path):
        raise ParsingError(f"File not found at path: {file_path}")

    lower_path = file_path.lower()

    if lower_path.endswith(".pdf"):
        return extract_pdf(file_path)
    elif lower_path.endswith(".docx"):
        return extract_docx(file_path)
    elif lower_path.endswith(".txt"):
        return extract_txt(file_path)
    else:
        raise ParsingError(f"Unsupported file format for path '{file_path}'. Supported: .pdf, .docx, .txt")


def extract_pdf(file_path: str) -> str:
    """
    Extract text and table contents from a PDF file using PyMuPDF.
    """
    text_blocks: List[str] = []

    try:
        doc = fitz.open(file_path)
        for page in doc:
            # 1. Extract regular text blocks
            page_text = page.get_text("text")
            if page_text and page_text.strip():
                text_blocks.append(page_text.strip())

            # 2. Extract tables safely
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    for table in tables.tables:
                        rows = table.extract()
                        for row in rows:
                            clean_row = [str(cell).strip() for cell in row if cell and str(cell).strip()]
                            if clean_row:
                                text_blocks.append(" | ".join(clean_row))
            except Exception as table_err:
                logger.warning("PDF table extraction warning on page %d of '%s': %s", page.number, file_path, table_err)

        doc.close()
    except Exception as exc:
        logger.exception("Failed to extract PDF text from '%s'.", file_path)
        raise ParsingError(f"PDF extraction failed: {exc}") from exc

    return "\n\n".join(text_blocks)


def extract_docx(file_path: str) -> str:
    """
    Extract text from paragraphs and tables in a DOCX file using python-docx.
    """
    text_blocks: List[str] = []

    try:
        doc = Document(file_path)

        # 1. Extract paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_blocks.append(para.text.strip())

        # 2. Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    # Remove duplicate adjacent cell texts caused by merged cells
                    unique_cells = []
                    for cell_t in row_cells:
                        if not unique_cells or unique_cells[-1] != cell_t:
                            unique_cells.append(cell_t)
                    text_blocks.append(" | ".join(unique_cells))

    except Exception as exc:
        logger.exception("Failed to extract DOCX text from '%s'.", file_path)
        raise ParsingError(f"DOCX extraction failed: {exc}") from exc

    return "\n\n".join(text_blocks)


def extract_txt(file_path: str) -> str:
    """
    Extract text from a plain TXT file with encoding fallbacks.
    """
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            raise ParsingError(f"TXT file read error: {exc}") from exc

    raise ParsingError(f"Could not decode TXT file '{file_path}' with standard encodings.")
