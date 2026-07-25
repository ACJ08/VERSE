import fitz  # PyMuPDF
from docx import Document


def extract_text(file_path: str) -> str:

    file_path = file_path.lower()

    if file_path.endswith(".pdf"):
        return extract_pdf(file_path)

    elif file_path.endswith(".docx"):
        return extract_docx(file_path)

    elif file_path.endswith(".txt"):
        return extract_txt(file_path)

    else:
        raise ValueError(
            "Unsupported file format"
        )



def extract_pdf(file_path: str) -> str:

    text = []

    doc = fitz.open(file_path)


    for page_number, page in enumerate(doc, start=1):

        # Extract normal text
        page_text = page.get_text("text")


        if page_text.strip():
            text.append(page_text)


        # Extract tables / blocks
        tables = page.find_tables()


        for table in tables.tables:

            rows = table.extract()


            for row in rows:

                clean_row = [
                    str(cell).strip()
                    for cell in row
                    if cell
                ]


                if clean_row:

                    text.append(
                        " | ".join(clean_row)
                    )


    doc.close()


    return "\n".join(text)




def extract_docx(file_path: str) -> str:

    doc = Document(file_path)

    text = []


    # Paragraph extraction
    for para in doc.paragraphs:

        if para.text.strip():

            text.append(
                para.text.strip()
            )



    # Table extraction
    for table in doc.tables:

        for row in table.rows:

            row_data = []


            for cell in row.cells:

                cell_text = cell.text.strip()


                if cell_text:

                    row_data.append(
                        cell_text
                    )


            if row_data:

                text.append(
                    " | ".join(row_data)
                )


    return "\n".join(text)




def extract_txt(file_path: str) -> str:

    with open(
        file_path,
        "r",
        encoding="utf-8"
    ) as f:

        return f.read()
        